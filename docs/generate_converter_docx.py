#!/usr/bin/env python3
"""Generate CONVERTER.docx from badapple/CONVERTER.md."""

import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.normpath(os.path.join(HERE, '..', 'badapple', 'CONVERTER.md'))
OUT = os.path.normpath(os.path.join(HERE, '..', 'badapple', 'CONVERTER.docx'))


def add_inline(p, text):
    """Render **bold** and `code` spans into paragraph p."""
    for tok in re.split(r'(\*\*.+?\*\*|`[^`]+`)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        else:
            p.add_run(tok)


def emit_table(doc, rows):
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=cols)
    tbl.style = 'Light Grid Accent 1'
    for ri, cells in enumerate(rows):
        cells = cells + [''] * (cols - len(cells))
        row = tbl.add_row()
        for ci, txt in enumerate(cells):
            cell = row.cells[ci]
            par = cell.paragraphs[0]
            add_inline(par, txt)
            if ri == 0:
                for r in par.runs:
                    r.bold = True
    doc.add_paragraph()


def main():
    doc = Document()
    # base font
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    in_code = False
    code_buf = []
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            emit_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run('\n'.join(code_buf))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            code_buf = []
            doc.add_paragraph()

    for line in lines:
        if line.startswith('```'):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # tables
        if line.lstrip().startswith('|'):
            if re.match(r'^\s*\|[\s:|-]+\|\s*$', line):
                continue  # separator
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            continue
        else:
            flush_table()

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip() == '---':
            continue
        elif not line.strip():
            continue
        else:
            text = line.strip()
            if text.startswith('> '):
                p = doc.add_paragraph(style='Intense Quote')
                add_inline(p, text[2:])
            elif text.startswith(('* ', '- ')):
                p = doc.add_paragraph(style='List Bullet')
                add_inline(p, text[2:])
            elif re.match(r'^\d+\.\s', text):
                p = doc.add_paragraph(style='List Number')
                add_inline(p, re.sub(r'^\d+\.\s', '', text))
            else:
                add_inline(doc.add_paragraph(), text)

    flush_table()
    doc.save(OUT)
    print('DOCX generated:', OUT)


if __name__ == '__main__':
    main()
