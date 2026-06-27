#!/usr/bin/env python3
"""
Build Bas/armcfgen.docx from Bas/armcfgen.md.

A small Markdown -> .docx converter (python-docx) covering the subset used by the
manual: # .. #### headings, fenced ``` code blocks, | pipe tables |, - bullets,
numbered items, > notes, and inline **bold** / `code`. Unicode passes through
untouched, and Word wraps table cells itself, so there is none of the latin-1 /
manual-wrapping fuss the PDF generator needs.

Convert to PDF afterwards with Word or LibreOffice (File > Export as PDF), or:
    soffice --headless --convert-to pdf Bas/armcfgen.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, '..', 'Bas', 'armcfgen.md')
OUT = os.path.join(HERE, '..', 'Bas', 'armcfgen.docx')

CODE_FONT = 'Consolas'
CODE_SIZE = 8.5
CODE_FILL = 'F2F2F2'
CODE_COLOR = RGBColor(0x33, 0x33, 0x33)


def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_inline(paragraph, text, base_bold=False):
    """Render a run sequence handling **bold** and `code`."""
    for tok in re.split(r'(\*\*.+?\*\*|`[^`]*`)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(9)
        else:
            r = paragraph.add_run(tok)
            if base_bold:
                r.bold = True


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    shade(p, CODE_FILL)
    first = True
    for ln in lines:
        r = p.add_run(ln if ln else ' ')
        r.font.name = CODE_FONT
        r.font.size = Pt(CODE_SIZE)
        r.font.color.rgb = CODE_COLOR
        if not first:
            pass
        first = False
        # line break between code lines (not after the last)
        if ln is not lines[-1]:
            r.add_break()


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = True
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].text = ''
        add_inline(cell.paragraphs[0], h, base_bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i in range(len(headers)):
            txt = row[i] if i < len(row) else ''
            cells[i].paragraphs[0].text = ''
            add_inline(cells[i].paragraphs[0], txt)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def convert(doc, md):
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith('```'):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i])
                i += 1
            i += 1
            if buf:
                add_code(doc, buf)
            continue

        # tables
        if stripped.startswith('|'):
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i].strip())
                i += 1
            rows = []
            for b in block:
                if re.match(r'^\|[\s:|-]+\|?$', b):   # separator
                    continue
                rows.append([c.strip() for c in b.strip('|').split('|')])
            if rows:
                add_table(doc, rows[0], rows[1:])
            continue

        # headings
        if line.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        elif line.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif line.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif line.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped == '---':
            pass  # section breaks are conveyed by headings
        elif stripped.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            add_inline(p, stripped[2:])
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            add_inline(p, stripped)
        elif stripped:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        i += 1


def main():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    # Force single line spacing - the default template line-spaces paragraphs,
    # which renders as "double spaced". Apply to every style we emit.
    for sname in ('Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3',
                  'List Bullet', 'Intense Quote'):
        try:
            pf = doc.styles[sname].paragraph_format
            pf.line_spacing = 1.0
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except KeyError:
            pass
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    with open(SRC, 'r', encoding='utf-8') as f:
        convert(doc, f.read())
    doc.save(OUT)
    print('Generated:', os.path.normpath(OUT))


if __name__ == '__main__':
    main()
